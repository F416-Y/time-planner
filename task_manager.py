#!/usr/bin/env python3
"""Time Planner - Safe JSON task manager for the time-planner skill.

Usage:
  python task_manager.py add --title "..." [options]
  python task_manager.py list [--status ...] [--priority ...]
  python task_manager.py modify <id> [options]
  python task_manager.py delete <id>
  python task_manager.py plan [--day-start 09:00] [--day-end 18:00] [--force] [--dynamic] [--quick]
  python task_manager.py review <id> <actual_minutes>
  python task_manager.py energy
  python task_manager.py limit [--set <hours>]
  python task_manager.py deps <id> [--add <dep_id>] [--remove <dep_id>]
  python task_manager.py config [--set-style strict|encourage|concise|playful|default]
  python task_manager.py predict <task_id>
  python task_manager.py warn
  python task_manager.py break <task_id>
  python task_manager.py stats
  python task_manager.py achievements
  python task_manager.py export <type> [--output <file.docx>] [--status ...] [--pdf]

All commands output JSON to stdout. Errors go to stderr.
"""

import argparse
import heapq
import json
import os
import sys
import uuid
from datetime import datetime, date, timedelta

SKILL_DIR = os.path.dirname(os.path.abspath(__file__))
CONFIG_FILE = os.path.join(SKILL_DIR, ".timeplanner_config.json")

_DEFAULT_TASKS_FILE = os.path.join(SKILL_DIR, "tasks.json")
_DEFAULT_ENERGY_FILE = os.path.join(SKILL_DIR, "energy_profile.json")
_DEFAULT_ACHIEVEMENTS_FILE = os.path.join(SKILL_DIR, "achievements.json")

_config_cache = None  # cached config dict, call load_config() to refresh


def load_config():
    """Load .timeplanner_config.json, return dict with defaults."""
    global _config_cache
    if _config_cache is not None:
        return _config_cache
    if not os.path.exists(CONFIG_FILE):
        _config_cache = {}
        return _config_cache
    try:
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            _config_cache = json.load(f)
    except (json.JSONDecodeError, IOError):
        _config_cache = {}
    if not isinstance(_config_cache, dict):
        _config_cache = {}
    return _config_cache


def save_config(config):
    """Persist config dict to .timeplanner_config.json."""
    global _config_cache
    _config_cache = config
    tmp = CONFIG_FILE + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(config, f, ensure_ascii=False, indent=2)
    os.replace(tmp, CONFIG_FILE)


def get_data_dir():
    """Return the configured data directory, or the skill dir if not set."""
    config = load_config()
    dd = config.get("data_dir")
    if dd and os.path.isdir(dd):
        return dd
    return SKILL_DIR


def _tasks_path():
    return os.path.join(get_data_dir(), "tasks.json")


def _energy_path():
    return os.path.join(get_data_dir(), "energy_profile.json")


def _achievements_path():
    return os.path.join(get_data_dir(), "achievements.json")

PRIORITY_WEIGHT = {"high": 3, "medium": 2, "low": 1}
VALID_STATUSES = {"pending", "in_progress", "completed", "cancelled"}
VALID_PRIORITIES = {"high", "medium", "low"}
VALID_STYLES = {"strict", "encourage", "concise", "playful", "default"}
VALID_ROUTINE_TYPES = {"daily", "weekly", "monthly"}
VALID_ROUTINE_STATUSES = {"active", "paused"}
ROUTINE_ID_PREFIX = "routine_"

# Break phases: (name, ratio, verb)
BREAK_PHASES = [
    ("调研", 0.20, "research"),
    ("整理", 0.30, "organize"),
    ("输出", 0.40, "output"),
    ("检查", 0.10, "review"),
]

ALL_ACHIEVEMENTS = [
    {"id": "FIRST_VICTORY", "name": "初次胜利", "description": "完成第一个任务", "icon": "TROPHY"},
    {"id": "STREAK_3", "name": "连续三日", "description": "连续3天有完成任务", "icon": "FIRE"},
    {"id": "BOSS_KILLER", "name": "Boss杀手", "description": "击败5个Boss任务", "icon": "SWORD"},
    {"id": "SHARPSHOOTER", "name": "神射手", "description": "复盘偏差率低于10%", "icon": "TARGET"},
    {"id": "WEEK_FULL", "name": "全勤一周", "description": "连续7天有完成任务", "icon": "STAR"},
    {"id": "OVERLOAD_REJECT", "name": "拒绝过载", "description": "累计3次触发过载保护", "icon": "SHIELD"},
    {"id": "DEVIATION_IMPROVED", "name": "偏差改善者", "description": "近7天平均偏差率下降超过20%", "icon": "CHART"},
]

ACHIEVEMENT_MAP = {a["id"]: a for a in ALL_ACHIEVEMENTS}


# ── task data ──────────────────────────────────────────────────────────────

def load_tasks():
    path = _tasks_path()
    if not os.path.exists(path):
        return []
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data.get("tasks", [])
    except (json.JSONDecodeError, IOError) as e:
        print(json.dumps({"error": f"tasks.json 文件损坏或无法读取: {e}"}, ensure_ascii=False), file=sys.stderr)
        sys.exit(1)


def save_tasks(tasks):
    path = _tasks_path()
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump({"tasks": tasks}, f, ensure_ascii=False, indent=2)
    os.replace(tmp, path)


def generate_id():
    return datetime.now().strftime("%Y%m%d") + "-" + uuid.uuid4().hex[:6]


def generate_routine_id():
    return ROUTINE_ID_PREFIX + datetime.now().strftime("%Y%m%d") + "-" + uuid.uuid4().hex[:6]


def is_routine_applicable(routine_task, target_date):
    """Check whether a routine task should appear on target_date.
    routine_days uses 1=Mon..7=Sun. Python weekday() returns 0=Mon..6=Sun.
    """
    rtype = routine_task.get("routine_type", "daily")
    if rtype == "daily":
        return True
    if rtype == "weekly":
        py_weekday = target_date.weekday()
        routine_weekday = py_weekday + 1
        return routine_weekday in (routine_task.get("routine_days") or [])
    if rtype == "monthly":
        return True
    return False


def get_applicable_routines(tasks, profile, target_date=None):
    """Return deviation-adjusted entries for all routines active today."""
    if target_date is None:
        target_date = date.today()

    result = []
    for t in tasks:
        if t.get("type") != "routine" or t.get("status") != "active":
            continue
        if not is_routine_applicable(t, target_date):
            continue

        raw_est = t.get("estimated_minutes", 30)
        adjusted = get_adjusted_minutes(t, profile)
        deviation = round((adjusted / raw_est - 1), 4) if raw_est > 0 else 0

        rt = t.get("routine_time") or "09:00"
        try:
            rt_h, rt_m = map(int, rt.split(":"))
        except (ValueError, AttributeError):
            rt_h, rt_m = 9, 0

        result.append({
            "task_id": t["id"],
            "title": t["title"],
            "priority": t.get("priority", "medium"),
            "estimated_minutes": raw_est,
            "adjusted_minutes": adjusted,
            "deviation_applied": deviation,
            "deadline": None,
            "tags": t.get("tags", []),
            "predicted_minutes": adjusted,
            "procrastination_risk": "low",
            "procrastination_score": 0,
            "is_routine": True,
            "routine_type": t.get("routine_type"),
            "routine_time": rt,
            "scheduled_start": rt_h * 60 + rt_m,
        })

    return result


def merge_routines_into_plan(normal_entries, routine_entries, day_start, day_end, dynamic=False):
    """Merge routine tasks into normal task plan, allocating sequential time blocks."""
    try:
        sh, sm = map(int, day_start.split(":"))
        eh, em = map(int, day_end.split(":"))
    except (ValueError, AttributeError):
        sh, sm = 9, 0
        eh, em = 18, 0
    day_start_min = sh * 60 + sm
    day_end_min = eh * 60 + em

    # assign preliminary cursor positions to normal tasks
    prelim_cursor = day_start_min
    normal_with_pos = []
    for at in normal_entries:
        normal_with_pos.append((prelim_cursor, at))
        prelim_cursor += at["adjusted_minutes"]

    # collect all items with sort keys (position, type_priority, entry)
    # type_priority: 0 = routine (prefer at same key), 1 = normal
    all_items = []
    for pos, at in normal_with_pos:
        all_items.append((pos, 1, at))
    for ra in routine_entries:
        all_items.append((ra["scheduled_start"], 0, ra))

    all_items.sort(key=lambda x: (x[0], x[1]))

    cursor = day_start_min
    plan = []
    next_best = None

    for _, _, entry in all_items:
        dur = entry["adjusted_minutes"]
        if cursor + dur > day_end_min:
            if dynamic and next_best is None:
                next_best = {
                    "task_id": entry["task_id"],
                    "title": entry["title"],
                    "priority": entry["priority"],
                    "adjusted_minutes": entry["adjusted_minutes"],
                }
            continue

        start_str = f"{cursor // 60:02d}:{cursor % 60:02d}"
        end_str = f"{(cursor + dur) // 60:02d}:{(cursor + dur) % 60:02d}"

        entry_out = dict(entry)
        entry_out["time"] = f"{start_str} - {end_str}"
        plan.append(entry_out)
        cursor += dur

    return plan, next_best


# ── energy profile data ────────────────────────────────────────────────────

def load_energy():
    path = _energy_path()
    if not os.path.exists(path):
        return {
            "daily_limit_hours": 8.0,
            "style": "default",
            "tag_profiles": {},
            "global_avg_deviation_rate": 0.0,
            "global_avg_procrastination": 0.0,
            "global_avg_deadline_adjusted": 0.0,
            "total_reviews": 0,
            "reviews": [],
        }
    try:
        with open(path, "r", encoding="utf-8") as f:
            profile = json.load(f)
    except (json.JSONDecodeError, IOError) as e:
        print(json.dumps({"error": f"energy_profile.json 文件损坏或无法读取: {e}"}, ensure_ascii=False), file=sys.stderr)
        sys.exit(1)
    if not isinstance(profile, dict):
        print(json.dumps({"error": "energy_profile.json 格式错误，顶层应为对象"}, ensure_ascii=False), file=sys.stderr)
        sys.exit(1)
    profile.setdefault("daily_limit_hours", 8.0)
    profile.setdefault("style", "default")
    profile.setdefault("global_avg_procrastination", 0.0)
    profile.setdefault("global_avg_deadline_adjusted", 0.0)
    # ensure every tag_profile has the new fields
    for tp in profile.get("tag_profiles", {}).values():
        tp.setdefault("avg_procrastination_count", 0.0)
        tp.setdefault("avg_deadline_adjusted", 0.0)
    return profile


def save_energy(profile):
    path = _energy_path()
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(profile, f, ensure_ascii=False, indent=2)
    os.replace(tmp, path)


def get_adjusted_minutes(task, profile):
    """Return estimated_minutes corrected by historical deviation rates.
    Priority: tag-specific avg > global avg > raw estimate. Capped at 120 min.
    """
    estimated = task.get("estimated_minutes", 30)
    tags = task.get("tags", [])

    if tags:
        rates = []
        for tag in tags:
            tp = profile.get("tag_profiles", {}).get(tag)
            if tp and tp.get("count", 0) > 0:
                rates.append(tp["avg_deviation_rate"])
        if rates:
            avg_rate = sum(rates) / len(rates)
            return min(max(5, round(estimated * (1 + avg_rate))), 120)

    global_rate = profile.get("global_avg_deviation_rate", 0.0)
    if global_rate:
        return min(max(5, round(estimated * (1 + global_rate))), 120)

    return estimated


# ── achievements data ──────────────────────────────────────────────────────

def load_achievements():
    path = _achievements_path()
    if not os.path.exists(path):
        return {
            "achievements": [],
            "stats": {
                "total_tasks_completed": 0,
                "total_bosses_defeated": 0,
                "current_streak": 0,
                "longest_streak": 0,
                "last_completed_date": None,
                "overload_trigger_count": 0,
            },
        }
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except (json.JSONDecodeError, IOError) as e:
        print(json.dumps({"error": f"achievements.json 文件损坏或无法读取: {e}"}, ensure_ascii=False), file=sys.stderr)
        sys.exit(1)
    if not isinstance(data, dict):
        print(json.dumps({"error": "achievements.json 格式错误，顶层应为对象"}, ensure_ascii=False), file=sys.stderr)
        sys.exit(1)
    data.setdefault("achievements", [])
    data.setdefault("stats", {})
    data["stats"].setdefault("total_tasks_completed", 0)
    data["stats"].setdefault("total_bosses_defeated", 0)
    data["stats"].setdefault("current_streak", 0)
    data["stats"].setdefault("longest_streak", 0)
    data["stats"].setdefault("last_completed_date", None)
    data["stats"].setdefault("overload_trigger_count", 0)
    return data


def save_achievements(data):
    path = _achievements_path()
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    os.replace(tmp, path)


def update_streak(ach_data, today_str):
    """Update streak counter based on today's completion."""
    stats = ach_data["stats"]
    last = stats["last_completed_date"]

    if last is None:
        stats["current_streak"] = 1
        stats["longest_streak"] = 1
    elif last == today_str:
        pass  # already counted today
    else:
        try:
            last_date = date.fromisoformat(last)
            today_date = date.fromisoformat(today_str)
            diff = (today_date - last_date).days
        except (ValueError, TypeError):
            diff = 999

        if diff == 1:
            stats["current_streak"] += 1
        elif diff > 1:
            stats["current_streak"] = 1
        # diff < 0 shouldn't happen, but reset if it does
        elif diff < 0:
            stats["current_streak"] = 1

        if stats["current_streak"] > stats["longest_streak"]:
            stats["longest_streak"] = stats["current_streak"]

    stats["last_completed_date"] = today_str


def check_achievements(ach_data, trigger, context):
    """Check for newly unlocked achievements. Returns list of newly unlocked."""
    existing_ids = {a["id"] for a in ach_data["achievements"]}
    stats = ach_data["stats"]
    profile = context.get("profile", {})
    new_achievements = []

    def unlock(ach_id):
        if ach_id not in existing_ids:
            entry = {
                **ACHIEVEMENT_MAP[ach_id],
                "unlocked_at": datetime.now().isoformat(),
            }
            ach_data["achievements"].append(entry)
            existing_ids.add(ach_id)
            new_achievements.append(entry)

    if trigger == "task_completed":
        if stats["total_tasks_completed"] >= 1:
            unlock("FIRST_VICTORY")
        if stats["current_streak"] >= 3:
            unlock("STREAK_3")
        if stats["current_streak"] >= 7:
            unlock("WEEK_FULL")

    if trigger == "boss_defeated":
        if stats["total_bosses_defeated"] >= 5:
            unlock("BOSS_KILLER")

    if trigger == "review":
        review = context.get("review", {})
        if abs(review.get("deviation_rate", 1.0)) < 0.10:
            unlock("SHARPSHOOTER")

        # deviation improved check
        reviews = profile.get("reviews", [])
        if len(reviews) >= 4:
            recent_7d, older_7d = _split_reviews_by_window(reviews, 7)
            if len(recent_7d) >= 2 and len(older_7d) >= 2:
                recent_avg = sum(r["deviation_rate"] for r in recent_7d) / len(recent_7d)
                older_avg = sum(r["deviation_rate"] for r in older_7d) / len(older_7d)
                if older_avg > 0 and (older_avg - recent_avg) / older_avg > 0.20:
                    unlock("DEVIATION_IMPROVED")

    if trigger == "overload":
        if stats.get("overload_trigger_count", 0) >= 3:
            unlock("OVERLOAD_REJECT")

    return new_achievements


def _split_reviews_by_window(reviews, days):
    """Split reviews into recent (within days) and older (days before that)."""
    now_ts = datetime.now().timestamp()
    cutoff = now_ts - days * 86400
    older_cutoff_start = cutoff - days * 86400

    recent = []
    older = []
    for r in reviews:
        try:
            rt = datetime.fromisoformat(r["reviewed_at"]).timestamp()
        except (ValueError, KeyError):
            continue
        if rt >= cutoff:
            recent.append(r)
        elif rt >= older_cutoff_start:
            older.append(r)

    return recent, older


# ── prediction helpers ─────────────────────────────────────────────────────

def compute_prediction(task, profile):
    """Compute predicted_minutes, confidence, and procrastination_risk for a task."""
    estimated = task.get("estimated_minutes", 30)
    adjusted = get_adjusted_minutes(task, profile)
    deviation_rate = round((adjusted / estimated - 1), 4) if estimated > 0 else 0.0

    tags = task.get("tags", [])
    tag_profiles = profile.get("tag_profiles", {})

    review_count = 0
    proc_counts = []
    for tag in tags:
        tp = tag_profiles.get(tag)
        if tp and tp.get("count", 0) > 0:
            review_count += tp["count"]
            proc_counts.append(tp.get("avg_procrastination_count", 0.0))

    if review_count >= 10:
        confidence = "high"
    elif review_count >= 3:
        confidence = "medium"
    else:
        confidence = "low"

    task_proc = task.get("procrastination_count", 0)
    if proc_counts:
        avg_proc = sum(proc_counts) / len(proc_counts)
    else:
        avg_proc = profile.get("global_avg_procrastination", 0.0)

    combined_risk = round(avg_proc * 0.6 + task_proc * 0.4, 1)

    if combined_risk >= 3:
        procrastination_risk = "high"
    elif combined_risk >= 1:
        procrastination_risk = "medium"
    else:
        procrastination_risk = "low"

    return {
        "predicted_minutes": adjusted,
        "deviation_rate": deviation_rate,
        "confidence": confidence,
        "procrastination_risk": procrastination_risk,
        "procrastination_score": combined_risk,
    }


# ── dependency graph ───────────────────────────────────────────────────────

def detect_cycle_path(tasks, task_id, dep_id):
    """Check if adding dep_id as dependency of task_id would create a cycle.
    Returns the cycle path (list of IDs) if found, None otherwise.
    """
    task_map = {t["id"]: t for t in tasks}

    def dfs(current, path, visited):
        if current == task_id:
            return [task_id] + path + [task_id]
        if current in visited:
            return None
        visited.add(current)
        task = task_map.get(current)
        if task:
            for d in task.get("dependencies", []):
                result = dfs(d, path + [current], visited)
                if result:
                    return result
        return None

    return dfs(dep_id, [], set())


def topological_sort(tasks):
    """Topological sort of tasks by their incomplete dependencies.
    Uses priority + deadline for tie-breaking within same dependency level.
    Returns (sorted_tasks, cycle_detected).
    """
    task_map = {t["id"]: t for t in tasks}
    task_ids = set(task_map.keys())

    in_degree = {tid: 0 for tid in task_ids}
    dependents = {tid: [] for tid in task_ids}

    for task in tasks:
        for dep_id in task.get("dependencies", []):
            if dep_id in task_ids:
                in_degree[task["id"]] += 1
                dependents[dep_id].append(task["id"])

    heap = []
    for tid, deg in in_degree.items():
        if deg == 0:
            t = task_map[tid]
            pw = PRIORITY_WEIGHT.get(t["priority"], 1)
            hd = 1 if t.get("deadline") else 0
            heapq.heappush(heap, (-pw, -hd, t.get("deadline") or "", tid))

    result = []
    while heap:
        _, _, _, tid = heapq.heappop(heap)
        result.append(task_map[tid])
        for dep_id in dependents[tid]:
            in_degree[dep_id] -= 1
            if in_degree[dep_id] == 0:
                t = task_map[dep_id]
                pw = PRIORITY_WEIGHT.get(t["priority"], 1)
                hd = 1 if t.get("deadline") else 0
                heapq.heappush(heap, (-pw, -hd, t.get("deadline") or "", dep_id))

    if len(result) != len(tasks):
        return tasks, True

    return result, False


def find_boss_tasks(incomplete_tasks):
    """Find 'Boss tasks' - incomplete blockers that hold up critical work.

    A blocker is a Boss task if it blocks at least one task that is:
      - priority=high, OR
      - has a deadline within 24 hours

    Returns dict: blocker_id -> {blocker, blocks, critical_count}
    """
    task_map = {t["id"]: t for t in incomplete_tasks}
    now = datetime.now()
    deadline_24h = now.timestamp() + 24 * 3600

    blocked_by = {}

    for task in incomplete_tasks:
        for dep_id in task.get("dependencies", []):
            dep_task = task_map.get(dep_id)
            if dep_task is None:
                continue
            is_critical = (
                task["priority"] == "high"
                or (
                    task.get("deadline")
                    and datetime.fromisoformat(task["deadline"]).timestamp() <= deadline_24h
                )
            )
            entry = {
                "task_id": task["id"],
                "title": task["title"],
                "priority": task["priority"],
                "deadline": task.get("deadline"),
                "is_critical": is_critical,
            }
            if dep_id not in blocked_by:
                blocked_by[dep_id] = []
            blocked_by[dep_id].append(entry)

    boss_tasks = {}
    for blocker_id, blocked_list in blocked_by.items():
        critical_count = sum(1 for b in blocked_list if b["is_critical"])
        if critical_count > 0:
            boss_tasks[blocker_id] = {
                "blocker": task_map[blocker_id],
                "blocks": blocked_list,
                "critical_count": critical_count,
            }

    return boss_tasks


# ── build adjusted task list ───────────────────────────────────────────────

def build_adjusted_task_list(tasks, profile, sort_by_priority=True):
    """Apply deviation correction to all tasks.

    Returns (adjusted_list, total_adjusted_minutes, adjustment_active).
    """
    if sort_by_priority:
        def sort_key(t):
            p = PRIORITY_WEIGHT.get(t["priority"], 1)
            has_deadline = 1 if t.get("deadline") else 0
            return (-p, -has_deadline, t.get("deadline") or "")
        tasks = sorted(tasks, key=sort_key)

    result = []
    total = 0
    adjustment_active = False

    for task in tasks:
        raw_est = task.get("estimated_minutes", 30)
        adjusted = get_adjusted_minutes(task, profile)
        deviation = round((adjusted / raw_est - 1), 4) if raw_est > 0 else 0
        if deviation != 0:
            adjustment_active = True

        prediction = compute_prediction(task, profile)
        entry = {
            "task_id": task["id"],
            "title": task["title"],
            "priority": task["priority"],
            "estimated_minutes": raw_est,
            "adjusted_minutes": adjusted,
            "deviation_applied": deviation,
            "deadline": task.get("deadline"),
            "tags": task.get("tags", []),
            "predicted_minutes": prediction["predicted_minutes"],
            "procrastination_risk": prediction["procrastination_risk"],
            "procrastination_score": prediction["procrastination_score"],
            "confidence": prediction["confidence"],
        }
        result.append(entry)
        total += adjusted

    return result, total, adjustment_active


def _annotate_boss_info(adjusted_entry, boss_tasks):
    """Annotate an adjusted_list entry with Boss task information."""
    tid = adjusted_entry["task_id"]

    if tid in boss_tasks:
        bt = boss_tasks[tid]
        adjusted_entry["is_boss"] = True
        adjusted_entry["boss_blocks_tasks"] = bt["blocks"]
        adjusted_entry["boss_critical_count"] = bt["critical_count"]

    for blocker_id, bt in boss_tasks.items():
        for blocked in bt["blocks"]:
            if blocked["task_id"] == tid:
                adjusted_entry["blocked_by_boss"] = blocker_id
                adjusted_entry["blocked_by_boss_title"] = bt["blocker"]["title"]
                break


# ── commands ────────────────────────────────────────────────────────────────

def cmd_add(args):
    tasks = load_tasks()
    now = datetime.now().isoformat()
    task = {
        "id": generate_id(),
        "title": args.title,
        "description": args.description or "",
        "priority": args.priority,
        "status": "pending",
        "estimated_minutes": args.estimated_minutes,
        "deadline": args.deadline,
        "created_at": now,
        "updated_at": now,
        "tags": [t.strip() for t in args.tags.split(",")] if args.tags else [],
    }
    tasks.append(task)
    save_tasks(tasks)
    print(json.dumps(task, ensure_ascii=False, indent=2))


def cmd_list(args):
    tasks = load_tasks()
    if args.status:
        tasks = [t for t in tasks if t["status"] == args.status]
    if args.priority:
        tasks = [t for t in tasks if t["priority"] == args.priority]

    def sort_key(t):
        done = 1 if t["status"] in ("completed", "cancelled") else 0
        p = PRIORITY_WEIGHT.get(t["priority"], 1)
        return (done, -p, t.get("deadline") or "")

    tasks.sort(key=sort_key)
    print(json.dumps(tasks, ensure_ascii=False, indent=2))


def cmd_modify(args):
    tasks = load_tasks()
    for t in tasks:
        if t["id"] == args.task_id:
            updates = {}
            deadline_changed = False

            if args.title is not None:
                updates["title"] = args.title
            if args.description is not None:
                updates["description"] = args.description
            if args.priority is not None:
                updates["priority"] = args.priority
            if args.status is not None:
                updates["status"] = args.status
            if args.estimated_minutes is not None:
                updates["estimated_minutes"] = args.estimated_minutes
            if args.deadline_flag:
                deadline_changed = True
                updates["deadline"] = None
            elif hasattr(args, "deadline"):
                deadline_changed = True
                updates["deadline"] = args.deadline
                current_deadline = t.get("deadline")
                if current_deadline and current_deadline != args.deadline:
                    updates["deadline_adjusted_count"] = t.get("deadline_adjusted_count", 0) + 1
            if args.tags is not None:
                updates["tags"] = [x.strip() for x in args.tags.split(",")] if args.tags else []

            if not updates:
                print(json.dumps({"error": "no fields to update"}, ensure_ascii=False))
                sys.exit(1)

            t.update(updates)
            t["updated_at"] = datetime.now().isoformat()
            save_tasks(tasks)
            print(json.dumps(t, ensure_ascii=False, indent=2))
            return
    print(json.dumps({"error": f"task {args.task_id} not found"}, ensure_ascii=False))
    sys.exit(1)


def cmd_delete(args):
    tasks = load_tasks()
    new_tasks = [t for t in tasks if t["id"] != args.task_id]
    if len(new_tasks) == len(tasks):
        print(json.dumps({"error": f"task {args.task_id} not found"}, ensure_ascii=False))
        sys.exit(1)
    # clean up dangling dependency references
    for t in new_tasks:
        deps = t.get("dependencies", [])
        if args.task_id in deps:
            t["dependencies"] = [d for d in deps if d != args.task_id]
            if not t["dependencies"]:
                del t["dependencies"]
            t["updated_at"] = datetime.now().isoformat()
    save_tasks(new_tasks)
    print(json.dumps({"deleted": args.task_id}, ensure_ascii=False))


def cmd_plan(args):
    tasks = load_tasks()
    profile = load_energy()
    # exclude parent tasks that have been split into sub-tasks, and routine tasks
    incomplete = [
        t for t in tasks
        if t["status"] not in ("completed", "cancelled")
        and t.get("type") != "routine"
        and not t.get("has_subtasks", False)
    ]

    sorted_incomplete, cycle_detected = topological_sort(incomplete)
    boss_tasks = find_boss_tasks(sorted_incomplete)

    adjusted_list, total_adjusted, adjustment_active = build_adjusted_task_list(
        sorted_incomplete, profile, sort_by_priority=False
    )

    for entry in adjusted_list:
        _annotate_boss_info(entry, boss_tasks)

    # ── routine tasks ──
    today = date.today()
    routine_adjusted = get_applicable_routines(tasks, profile, today)
    if routine_adjusted:
        adjustment_active = True
    routine_minutes = sum(r["adjusted_minutes"] for r in routine_adjusted)

    daily_limit_hours = profile.get("daily_limit_hours", 8.0)
    daily_limit_minutes = int(daily_limit_hours * 60)

    # ── overload detection (routine tasks excluded from limit) ──
    if total_adjusted > daily_limit_minutes and not args.force:
        feasible = []
        overload = []
        cursor = 0
        for at in adjusted_list:
            if cursor + at["adjusted_minutes"] <= daily_limit_minutes:
                feasible.append(at)
                cursor += at["adjusted_minutes"]
            else:
                overload.append({**at, "reason": "超出每日上限"})

        # increment procrastination_count for overloaded tasks
        for ol in overload:
            tid = ol["task_id"]
            for t in tasks:
                if t["id"] == tid:
                    t["procrastination_count"] = t.get("procrastination_count", 0) + 1
                    t["updated_at"] = datetime.now().isoformat()
                    break
        save_tasks(tasks)

        # track overload trigger count and check achievements
        ach_data = load_achievements()
        ach_data["stats"]["overload_trigger_count"] = ach_data["stats"].get("overload_trigger_count", 0) + 1
        new_achs = check_achievements(ach_data, "overload", {"profile": profile, "tasks": tasks})
        save_achievements(ach_data)

        # merge routines into feasible plan
        combined_plan, plan_next_best = merge_routines_into_plan(
            feasible, routine_adjusted, args.day_start, args.day_end, args.dynamic
        )

        # ── quick mode: auto-proceed with feasible plan, no confirmation ──
        if args.quick:
            normal_count = sum(1 for e in combined_plan if not e.get("is_routine"))
            quick_output = {
                "day_range": f"{args.day_start} - {args.day_end}",
                "planned_count": len(combined_plan),
                "remaining_tasks": len(adjusted_list) - normal_count,
                "routine_count": sum(1 for e in combined_plan if e.get("is_routine")),
                "energy_adjustment_active": adjustment_active,
                "plan": combined_plan,
                "quick": True,
                "overload_skipped": {
                    "total_adjusted_minutes": total_adjusted,
                    "daily_limit_minutes": daily_limit_minutes,
                    "daily_limit_hours": daily_limit_hours,
                    "overload_count": len(overload),
                    "overload_tasks": overload,
                },
            }
            if args.dynamic and overload:
                quick_output["next_best"] = overload[0]["task_id"]
            if cycle_detected:
                quick_output["cycle_warning"] = "检测到循环依赖，已降级为原始排序，建议手动修复 tasks.json。"
            if boss_tasks:
                quick_output["boss_tasks_detected"] = True
                quick_output["boss_task_ids"] = list(boss_tasks.keys())
            if new_achs:
                quick_output["new_achievements"] = new_achs
            print(json.dumps(quick_output, ensure_ascii=False, indent=2))
            return

        suggestion = (
            f"今日修正后总耗时 {total_adjusted} 分钟（{total_adjusted / 60:.1f} 小时），"
            f"超过每日上限 {daily_limit_minutes} 分钟（{daily_limit_hours} 小时）。"
            f"已按优先级为你筛选 {len(feasible)} 项可行任务（共 {cursor} 分钟），"
            f"{len(overload)} 项建议延迟到明天或拆分。"
            f"另有 {len(routine_adjusted)} 项例行任务（{routine_minutes} 分钟）不受过载限制。"
            f"输入「确认」采纳此计划并生成时间块，或指定要调整的任务。"
        )

        output = {
            "warning": "overload",
            "total_adjusted_minutes": total_adjusted,
            "daily_limit_minutes": daily_limit_minutes,
            "daily_limit_hours": daily_limit_hours,
            "feasible_count": len(feasible),
            "overload_count": len(overload),
            "routine_count": len(routine_adjusted),
            "routine_minutes": routine_minutes,
            "feasible_plan": combined_plan,
            "overload_tasks": overload,
            "suggestion": suggestion,
        }
        if args.dynamic and overload:
            output["next_best"] = overload[0]["task_id"]
        if cycle_detected:
            output["cycle_warning"] = "检测到循环依赖，已降级为原始排序，建议手动修复 tasks.json。"
        if boss_tasks:
            output["boss_tasks_detected"] = True
            output["boss_task_ids"] = list(boss_tasks.keys())
        if new_achs:
            output["new_achievements"] = new_achs
        print(json.dumps(output, ensure_ascii=False, indent=2))
        return

    # ── normal plan (or forced) with routines merged ──
    combined_plan, next_best = merge_routines_into_plan(
        adjusted_list, routine_adjusted, args.day_start, args.day_end, args.dynamic
    )

    normal_count = sum(1 for e in combined_plan if not e.get("is_routine"))

    output = {
        "day_range": f"{args.day_start} - {args.day_end}",
        "planned_count": len(combined_plan),
        "remaining_tasks": len(adjusted_list) - normal_count,
        "routine_count": sum(1 for e in combined_plan if e.get("is_routine")),
        "energy_adjustment_active": adjustment_active,
        "plan": combined_plan,
    }
    if args.force:
        output["forced"] = True
    if args.quick:
        output["quick"] = True
    if args.dynamic and next_best:
        output["next_best"] = next_best
    if cycle_detected:
        output["cycle_warning"] = "检测到循环依赖，已降级为原始排序，建议手动修复 tasks.json。"
    if boss_tasks:
        output["boss_tasks_detected"] = True
        output["boss_task_ids"] = list(boss_tasks.keys())
    print(json.dumps(output, ensure_ascii=False, indent=2))


def cmd_review(args):
    tasks = load_tasks()
    task = None
    for t in tasks:
        if t["id"] == args.task_id:
            task = t
            break
    if not task:
        print(json.dumps({"error": f"task {args.task_id} not found"}, ensure_ascii=False))
        sys.exit(1)

    estimated = task.get("estimated_minutes", 30)
    actual = args.actual_minutes
    if actual <= 0:
        print(json.dumps({"error": "actual_minutes must be > 0"}, ensure_ascii=False))
        sys.exit(1)

    deviation_minutes = actual - estimated
    deviation_rate = round((actual - estimated) / estimated, 4)

    profile = load_energy()

    review = {
        "id": "rev-" + generate_id(),
        "task_id": task["id"],
        "title": task["title"],
        "tags": task.get("tags", []),
        "estimated_minutes": estimated,
        "actual_minutes": actual,
        "deviation_minutes": deviation_minutes,
        "deviation_rate": deviation_rate,
        "reviewed_at": datetime.now().isoformat(),
    }
    profile["reviews"].append(review)
    profile["total_reviews"] = len(profile["reviews"])

    tags = task.get("tags", [])
    if not tags:
        tags = ["_untagged"]

    tag_impacts = {}
    for tag in tags:
        if tag not in profile["tag_profiles"]:
            profile["tag_profiles"][tag] = {
                "review_ids": [],
                "avg_deviation_rate": 0.0,
                "count": 0,
                "avg_procrastination_count": 0.0,
                "avg_deadline_adjusted": 0.0,
            }
        tp = profile["tag_profiles"][tag]
        tp.setdefault("avg_procrastination_count", 0.0)
        tp.setdefault("avg_deadline_adjusted", 0.0)
        tp["review_ids"].append(review["id"])
        tp["count"] = len(tp["review_ids"])
        rates = [
            r["deviation_rate"]
            for r in profile["reviews"]
            if r["id"] in tp["review_ids"]
        ]
        tp["avg_deviation_rate"] = round(sum(rates) / len(rates), 4)

        # also update procrastination/deadline_adjusted averages from the task
        proc_count = task.get("procrastination_count", 0)
        dl_adj_count = task.get("deadline_adjusted_count", 0)
        # rolling average update
        old_proc = tp["avg_procrastination_count"]
        old_dl = tp["avg_deadline_adjusted"]
        tp["avg_procrastination_count"] = round(
            (old_proc * (tp["count"] - 1) + proc_count) / tp["count"], 4
        )
        tp["avg_deadline_adjusted"] = round(
            (old_dl * (tp["count"] - 1) + dl_adj_count) / tp["count"], 4
        )

        tag_impacts[tag] = tp["avg_deviation_rate"]

    if profile["reviews"]:
        all_rates = [r["deviation_rate"] for r in profile["reviews"]]
        profile["global_avg_deviation_rate"] = round(sum(all_rates) / len(all_rates), 4)
        # update global procrastination average
        all_tasks = load_tasks()
        proc_vals = [t.get("procrastination_count", 0) for t in all_tasks]
        dl_adj_vals = [t.get("deadline_adjusted_count", 0) for t in all_tasks]
        profile["global_avg_procrastination"] = round(sum(proc_vals) / max(len(proc_vals), 1), 4)
        profile["global_avg_deadline_adjusted"] = round(sum(dl_adj_vals) / max(len(dl_adj_vals), 1), 4)

    save_energy(profile)

    # detect Boss tasks BEFORE status change
    all_incomplete_before = [t for t in tasks if t["status"] not in ("completed", "cancelled")]
    boss_tasks_before = find_boss_tasks(all_incomplete_before)
    was_boss_before = task["id"] in boss_tasks_before

    # guard against reviewing a parent task that has been split
    if task.get("has_subtasks"):
        print(json.dumps({"error": "该任务已拆分为子任务，请对子任务分别复盘"}, ensure_ascii=False))
        sys.exit(1)

    # auto-complete task if not already completed
    was_completed = task["status"] in ("completed", "cancelled")
    if not was_completed:
        task["status"] = "completed"
        task["updated_at"] = datetime.now().isoformat()
        is_parent = task.get("has_subtasks", False)
        is_sub = task.get("parent_task_id") is not None

        # check if all sub-tasks of parent are done → auto-complete parent
        parent_auto_completed = None
        if is_sub:
            parent_id = task["parent_task_id"]
            parent_task = None
            for t in tasks:
                if t["id"] == parent_id:
                    parent_task = t
                    break
            if parent_task:
                siblings = [t for t in tasks if t.get("parent_task_id") == parent_id]
                all_done = all(s["status"] in ("completed", "cancelled") for s in siblings)
                if all_done and parent_task["status"] not in ("completed", "cancelled"):
                    parent_task["status"] = "completed"
                    parent_task["updated_at"] = datetime.now().isoformat()
                    parent_auto_completed = {
                        "task_id": parent_task["id"],
                        "title": parent_task["title"],
                    }

        save_tasks(tasks)

        # ── achievements & streak ──
        ach_data = load_achievements()
        stats = ach_data["stats"]
        stats["total_tasks_completed"] += 1

        today_str = date.today().isoformat()
        update_streak(ach_data, today_str)

        if was_boss_before:
            stats["total_bosses_defeated"] += 1

        context = {
            "task": task,
            "review": review,
            "profile": profile,
            "tasks": tasks,
        }
        new_achs = check_achievements(ach_data, "task_completed", context)
        boss_achs = check_achievements(ach_data, "boss_defeated", context)
        review_achs = check_achievements(ach_data, "review", context)
        new_achs = new_achs + [a for a in boss_achs if a not in new_achs] + [a for a in review_achs if a not in new_achs]
        save_achievements(ach_data)

        sign = "+" if deviation_minutes > 0 else ""
        result = {
            "review": review,
            "deviation_readable": f"{sign}{deviation_minutes} min ({sign}{round(deviation_rate * 100, 1)}%)",
            "tag_impacts": tag_impacts,
            "task_auto_completed": True,
        }
        if parent_auto_completed:
            result["parent_auto_completed"] = parent_auto_completed
        if new_achs:
            result["new_achievements"] = new_achs
        if was_boss_before:
            result["boss_defeated"] = True
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return

    save_tasks(tasks)
    sign = "+" if deviation_minutes > 0 else ""
    print(json.dumps({
        "review": review,
        "deviation_readable": f"{sign}{deviation_minutes} min ({sign}{round(deviation_rate * 100, 1)}%)",
        "tag_impacts": tag_impacts,
        "task_already_completed": True,
    }, ensure_ascii=False, indent=2))


def cmd_energy(args):
    profile = load_energy()

    tag_summary = {}
    for tag, tp in profile.get("tag_profiles", {}).items():
        sign = "+" if tp["avg_deviation_rate"] > 0 else ""
        tag_summary[tag] = {
            "count": tp["count"],
            "avg_deviation_rate": tp["avg_deviation_rate"],
            "readable": f"{sign}{round(tp['avg_deviation_rate'] * 100, 1)}%",
            "avg_procrastination": tp.get("avg_procrastination_count", 0.0),
            "avg_deadline_adjusted": tp.get("avg_deadline_adjusted", 0.0),
        }

    global_rate = profile.get("global_avg_deviation_rate", 0.0)
    global_sign = "+" if global_rate > 0 else ""

    print(json.dumps({
        "daily_limit_hours": profile.get("daily_limit_hours", 8.0),
        "style": profile.get("style", "default"),
        "total_reviews": profile.get("total_reviews", 0),
        "global_avg_deviation": f"{global_sign}{round(global_rate * 100, 1)}%",
        "global_avg_deviation_rate": global_rate,
        "global_avg_procrastination": profile.get("global_avg_procrastination", 0.0),
        "global_avg_deadline_adjusted": profile.get("global_avg_deadline_adjusted", 0.0),
        "tag_profiles": tag_summary,
    }, ensure_ascii=False, indent=2))


def cmd_limit(args):
    profile = load_energy()

    if args.set is not None:
        if args.set <= 0:
            print(json.dumps({"error": "daily limit must be > 0 hours"}, ensure_ascii=False))
            sys.exit(1)
        old_limit = profile.get("daily_limit_hours", 8.0)
        profile["daily_limit_hours"] = args.set
        save_energy(profile)
        print(json.dumps({
            "daily_limit_hours": args.set,
            "previous": old_limit,
            "updated": True,
        }, ensure_ascii=False, indent=2))
    else:
        print(json.dumps({
            "daily_limit_hours": profile.get("daily_limit_hours", 8.0),
            "daily_limit_minutes": int(profile.get("daily_limit_hours", 8.0) * 60),
        }, ensure_ascii=False, indent=2))


def cmd_deps(args):
    """View / add / remove task dependencies."""
    tasks = load_tasks()
    task_map = {t["id"]: t for t in tasks}

    target = task_map.get(args.task_id)
    if not target:
        print(json.dumps({"error": f"task {args.task_id} not found"}, ensure_ascii=False))
        sys.exit(1)

    if not args.add and not args.remove:
        deps = []
        for dep_id in target.get("dependencies", []):
            dep = task_map.get(dep_id)
            deps.append({
                "task_id": dep_id,
                "title": dep["title"] if dep else "(已删除)",
                "status": dep["status"] if dep else "unknown",
            })

        dependents = []
        for t in tasks:
            if args.task_id in t.get("dependencies", []):
                dependents.append({
                    "task_id": t["id"],
                    "title": t["title"],
                    "status": t["status"],
                    "priority": t["priority"],
                })

        print(json.dumps({
            "task_id": args.task_id,
            "title": target["title"],
            "dependencies": deps,
            "dependents": dependents,
            "dependency_count": len(deps),
            "dependent_count": len(dependents),
        }, ensure_ascii=False, indent=2))
        return

    if args.add:
        dep_id = args.add
        if dep_id not in task_map:
            print(json.dumps({"error": f"dependency task {dep_id} not found"}, ensure_ascii=False))
            sys.exit(1)
        if dep_id == args.task_id:
            print(json.dumps({"error": "任务不能依赖自己"}, ensure_ascii=False))
            sys.exit(1)

        deps = target.get("dependencies", [])
        if dep_id in deps:
            print(json.dumps({"error": f"依赖 {dep_id} 已存在"}, ensure_ascii=False))
            sys.exit(1)

        cycle = detect_cycle_path(tasks, args.task_id, dep_id)
        if cycle:
            path_str = " -> ".join(cycle)
            print(json.dumps({
                "error": f"[BOSS-CYCLE] 检测到循环依赖：{path_str} -> {args.task_id}，已拒绝操作",
                "cycle_path": cycle,
            }, ensure_ascii=False))
            sys.exit(1)

        target.setdefault("dependencies", [])
        target["dependencies"].append(dep_id)
        target["updated_at"] = datetime.now().isoformat()
        save_tasks(tasks)
        print(json.dumps({
            "added_dependency": {"from": args.task_id, "to": dep_id},
            "task": target,
        }, ensure_ascii=False, indent=2))
        return

    if args.remove:
        dep_id = args.remove
        deps = target.get("dependencies", [])
        if dep_id not in deps:
            print(json.dumps({"error": f"依赖 {dep_id} 不存在于任务 {args.task_id}"}, ensure_ascii=False))
            sys.exit(1)

        target["dependencies"] = [d for d in deps if d != dep_id]
        if not target["dependencies"]:
            del target["dependencies"]
        target["updated_at"] = datetime.now().isoformat()
        save_tasks(tasks)
        print(json.dumps({
            "removed_dependency": {"from": args.task_id, "to": dep_id},
            "task": target,
        }, ensure_ascii=False, indent=2))
        return


def cmd_config(args):
    """Show or set config (style, data_dir, etc.) stored in .timeplanner_config.json and energy_profile.json."""
    profile = load_energy()
    config = load_config()

    if args.set_style is not None:
        style = args.set_style
        if style not in VALID_STYLES:
            print(json.dumps({
                "error": f"无效的风格值 '{style}'，可选值: {sorted(VALID_STYLES)}",
            }, ensure_ascii=False))
            sys.exit(1)
        old_style = profile.get("style", "default")
        profile["style"] = style
        save_energy(profile)
        print(json.dumps({
            "style": style,
            "previous": old_style,
            "updated": True,
        }, ensure_ascii=False, indent=2))
    elif args.set_data_dir is not None:
        new_dir = args.set_data_dir
        if not os.path.isdir(new_dir):
            print(json.dumps({
                "error": f"目录不存在或无法访问: {new_dir}",
            }, ensure_ascii=False))
            sys.exit(1)
        old_dir = config.get("data_dir") or SKILL_DIR
        config["data_dir"] = new_dir
        save_config(config)
        # force reload on next access
        load_config()
        print(json.dumps({
            "data_dir": new_dir,
            "previous": old_dir,
            "updated": True,
            "note": "数据文件将从新目录读写。如果新目录中没有 tasks.json 等文件，将自动创建。",
        }, ensure_ascii=False, indent=2))
    elif args.show_data_dir:
        data_dir = get_data_dir()
        print(json.dumps({
            "data_dir": data_dir,
            "is_custom": "data_dir" in config,
            "default_dir": SKILL_DIR,
        }, ensure_ascii=False, indent=2))
    elif args.reset_data_dir:
        if "data_dir" in config:
            old_dir = config.pop("data_dir")
            save_config(config)
            # force reload
            load_config()
            print(json.dumps({
                "data_dir": SKILL_DIR,
                "previous": old_dir,
                "reset": True,
                "note": "已恢复默认数据目录",
            }, ensure_ascii=False, indent=2))
        else:
            print(json.dumps({
                "data_dir": SKILL_DIR,
                "note": "当前已是默认数据目录，无需重置",
            }, ensure_ascii=False, indent=2))
    else:
        data_dir = get_data_dir()
        print(json.dumps({
            "style": profile.get("style", "default"),
            "daily_limit_hours": profile.get("daily_limit_hours", 8.0),
            "total_reviews": profile.get("total_reviews", 0),
            "data_dir": data_dir,
            "is_custom_data_dir": "data_dir" in config,
            "default_dir": SKILL_DIR,
        }, ensure_ascii=False, indent=2))


# ── Module 十: predict ─────────────────────────────────────────────────────

def cmd_predict(args):
    """Predict task duration and procrastination risk."""
    tasks = load_tasks()
    task = None
    for t in tasks:
        if t["id"] == args.task_id:
            task = t
            break
    if not task:
        print(json.dumps({"error": f"task {args.task_id} not found"}, ensure_ascii=False))
        sys.exit(1)

    profile = load_energy()
    prediction = compute_prediction(task, profile)

    tags = task.get("tags", [])
    tag_profiles = profile.get("tag_profiles", {})
    tag_details = {}
    for tag in tags:
        tp = tag_profiles.get(tag)
        if tp and tp.get("count", 0) > 0:
            tag_details[tag] = {
                "count": tp["count"],
                "avg_deviation_rate": tp["avg_deviation_rate"],
                "avg_procrastination": tp.get("avg_procrastination_count", 0.0),
                "avg_deadline_adjusted": tp.get("avg_deadline_adjusted", 0.0),
            }

    output = {
        "task_id": task["id"],
        "title": task["title"],
        "estimated_minutes": task.get("estimated_minutes", 30),
        "predicted_minutes": prediction["predicted_minutes"],
        "deviation_rate": prediction["deviation_rate"],
        "confidence": prediction["confidence"],
        "procrastination_risk": prediction["procrastination_risk"],
        "procrastination_score": prediction["procrastination_score"],
        "tag_details": tag_details,
    }

    if prediction["procrastination_risk"] == "high":
        output["procrastination_warning"] = "该任务拖延风险较高，建议优先安排并设置明确的截止时间"

    print(json.dumps(output, ensure_ascii=False, indent=2))


# ── Module 十: warn ─────────────────────────────────────────────────────────

def cmd_warn(args):
    """Scan all incomplete tasks for high procrastination risk."""
    tasks = load_tasks()
    profile = load_energy()
    incomplete = [t for t in tasks if t["status"] not in ("completed", "cancelled")]

    warnings = []
    for task in incomplete:
        prediction = compute_prediction(task, profile)
        if prediction["procrastination_risk"] == "high":
            warnings.append({
                "task_id": task["id"],
                "title": task["title"],
                "priority": task["priority"],
                "deadline": task.get("deadline"),
                "procrastination_score": prediction["procrastination_score"],
                "procrastination_risk": prediction["procrastination_risk"],
                "procrastination_count": task.get("procrastination_count", 0),
                "deadline_adjusted_count": task.get("deadline_adjusted_count", 0),
            })

    warnings.sort(key=lambda w: w["procrastination_score"], reverse=True)

    print(json.dumps({
        "warning_count": len(warnings),
        "scanned_count": len(incomplete),
        "warnings": warnings,
    }, ensure_ascii=False, indent=2))


# ── Module 十一: break ─────────────────────────────────────────────────────

def cmd_break(args):
    """Auto-split a large task into sub-tasks with verb-guided phases."""
    tasks = load_tasks()
    task = None
    for t in tasks:
        if t["id"] == args.task_id:
            task = t
            break
    if not task:
        print(json.dumps({"error": f"task {args.task_id} not found"}, ensure_ascii=False))
        sys.exit(1)

    estimated = task.get("estimated_minutes", 30)
    if estimated <= 120:
        print(json.dumps({
            "error": f"任务预估 {estimated} 分钟 <= 120 分钟，无需拆分。只有超过 120 分钟的任务才能使用 break 命令。",
            "estimated_minutes": estimated,
        }, ensure_ascii=False))
        sys.exit(1)

    if task.get("has_subtasks"):
        print(json.dumps({"error": "该任务已经拆分过，请勿重复操作"}, ensure_ascii=False))
        sys.exit(1)

    now = datetime.now().isoformat()
    sub_tasks = []
    prev_id = None

    for i, (phase_name, ratio, _verb) in enumerate(BREAK_PHASES, 1):
        sub_id = f"{task['id']}-sub-{i}"
        sub_minutes = max(5, round(estimated * ratio))
        sub_task = {
            "id": sub_id,
            "title": f"[{phase_name}] {task['title']}",
            "description": f"父任务「{task['title']}」的{phase_name}阶段（自动拆分 第{i}/4步）",
            "priority": task["priority"],
            "status": "pending",
            "estimated_minutes": sub_minutes,
            "deadline": task.get("deadline"),
            "created_at": now,
            "updated_at": now,
            "tags": task.get("tags", []),
            "parent_task_id": task["id"],
        }
        if prev_id:
            sub_task["dependencies"] = [prev_id]
        sub_tasks.append(sub_task)
        prev_id = sub_id

    task["has_subtasks"] = True
    task["updated_at"] = now
    tasks.extend(sub_tasks)

    # re-point dependents from parent to last sub-task
    last_sub_id = sub_tasks[-1]["id"]
    for t in tasks:
        deps = t.get("dependencies", [])
        if task["id"] in deps:
            new_deps = [last_sub_id if d == task["id"] else d for d in deps]
            t["dependencies"] = new_deps
            t["updated_at"] = now

    save_tasks(tasks)
    print(json.dumps({
        "parent_task_id": task["id"],
        "parent_title": task["title"],
        "original_minutes": estimated,
        "sub_tasks": sub_tasks,
        "note": "后续任务依赖已自动指向最后一个子任务",
    }, ensure_ascii=False, indent=2))


# ── Module 十二: stats ─────────────────────────────────────────────────────

def cmd_stats(args):
    """Show growth tracking statistics."""
    ach_data = load_achievements()
    stats = ach_data["stats"]

    print(json.dumps({
        "total_tasks_completed": stats.get("total_tasks_completed", 0),
        "total_bosses_defeated": stats.get("total_bosses_defeated", 0),
        "current_streak": stats.get("current_streak", 0),
        "longest_streak": stats.get("longest_streak", 0),
        "last_completed_date": stats.get("last_completed_date"),
        "overload_trigger_count": stats.get("overload_trigger_count", 0),
        "total_achievements": len(ach_data["achievements"]),
        "achievement_ids": [a["id"] for a in ach_data["achievements"]],
    }, ensure_ascii=False, indent=2))


# ── Module 十二: achievements ──────────────────────────────────────────────

def cmd_achievements(args):
    """List all achievements with unlock status."""
    ach_data = load_achievements()
    unlocked_ids = {a["id"] for a in ach_data["achievements"]}
    unlocked_map = {a["id"]: a for a in ach_data["achievements"]}

    all_with_status = []
    for ach_def in ALL_ACHIEVEMENTS:
        ach_id = ach_def["id"]
        if ach_id in unlocked_ids:
            entry = {
                **ach_def,
                "unlocked": True,
                "unlocked_at": unlocked_map[ach_id].get("unlocked_at"),
            }
        else:
            entry = {**ach_def, "unlocked": False}
        all_with_status.append(entry)

    print(json.dumps({
        "achievements": all_with_status,
        "unlocked_count": len(unlocked_ids),
        "total_count": len(ALL_ACHIEVEMENTS),
        "stats": ach_data["stats"],
    }, ensure_ascii=False, indent=2))


# ── Module 十四: routine ────────────────────────────────────────────────────

def cmd_routine(args):
    """Dispatch routine sub-actions: add, list, done, log, pause, resume."""
    action = args.action
    if action == "add":
        cmd_routine_add(args)
    elif action == "list":
        cmd_routine_list(args)
    elif action == "done":
        cmd_routine_done(args)
    elif action == "log":
        cmd_routine_log(args)
    elif action == "pause":
        cmd_routine_pause(args)
    elif action == "resume":
        cmd_routine_resume(args)


def cmd_routine_add(args):
    tasks = load_tasks()
    now = datetime.now().isoformat()

    if args.routine_type not in VALID_ROUTINE_TYPES:
        print(json.dumps({"error": f"routine_type 必须是 {sorted(VALID_ROUTINE_TYPES)} 之一"}, ensure_ascii=False))
        sys.exit(1)

    routine_days = None
    if args.routine_days:
        try:
            routine_days = [int(x.strip()) for x in args.routine_days.split(",")]
            for d in routine_days:
                if d < 1 or d > 7:
                    raise ValueError(f"星期几必须在 1-7 之间，收到: {d}")
        except ValueError as e:
            print(json.dumps({"error": f"routine_days 格式错误: {e}。应为逗号分隔的数字 1-7 (1=周一)"}, ensure_ascii=False))
            sys.exit(1)

    if args.routine_type == "weekly" and not routine_days:
        print(json.dumps({"error": "weekly 类型必须指定 --routine-days (逗号分隔的数字 1-7，1=周一)"}, ensure_ascii=False))
        sys.exit(1)

    task = {
        "id": generate_routine_id(),
        "title": args.title,
        "type": "routine",
        "routine_type": args.routine_type,
        "routine_time": args.routine_time or "09:00",
        "description": args.description or "",
        "priority": args.priority or "medium",
        "status": "active",
        "estimated_minutes": args.estimated_minutes or 30,
        "deadline": None,
        "created_at": now,
        "updated_at": now,
        "tags": [t.strip() for t in args.tags.split(",")] if args.tags else [],
        "completion_log": [],
    }
    if routine_days is not None:
        task["routine_days"] = routine_days

    tasks.append(task)
    save_tasks(tasks)
    print(json.dumps(task, ensure_ascii=False, indent=2))


def cmd_routine_list(args):
    tasks = load_tasks()
    routines = [t for t in tasks if t.get("type") == "routine"]

    today = date.today()
    monday = today - timedelta(days=today.weekday())
    first_of_month = today.replace(day=1)

    result = []
    for r in routines:
        comp_log = r.get("completion_log") or []
        this_week = sum(1 for c in comp_log if date.fromisoformat(c["date"]) >= monday)
        this_month = sum(1 for c in comp_log if date.fromisoformat(c["date"]) >= first_of_month)

        result.append({
            "id": r["id"],
            "title": r["title"],
            "routine_type": r.get("routine_type"),
            "routine_time": r.get("routine_time"),
            "routine_days": r.get("routine_days"),
            "status": r.get("status"),
            "estimated_minutes": r.get("estimated_minutes", 30),
            "tags": r.get("tags", []),
            "this_week_count": this_week,
            "this_month_count": this_month,
            "total_count": len(comp_log),
        })

    print(json.dumps(result, ensure_ascii=False, indent=2))


def cmd_routine_done(args):
    tasks = load_tasks()
    task = None
    for t in tasks:
        if t["id"] == args.task_id and t.get("type") == "routine":
            task = t
            break

    if not task:
        print(json.dumps({"error": f"例行任务 {args.task_id} 未找到"}, ensure_ascii=False))
        sys.exit(1)

    if task.get("status") != "active":
        print(json.dumps({"error": "只能标记 active 状态的例行任务为完成"}, ensure_ascii=False))
        sys.exit(1)

    now = datetime.now()
    today_str = now.strftime("%Y-%m-%d")
    time_str = now.strftime("%H:%M")
    entry = {"date": today_str, "time": time_str}

    task.setdefault("completion_log", []).append(entry)
    task["updated_at"] = now.isoformat()
    save_tasks(tasks)

    print(json.dumps({
        "task_id": task["id"],
        "title": task["title"],
        "completed_at": f"{today_str} {time_str}",
        "completion_log": task["completion_log"],
        "total_completions": len(task["completion_log"]),
    }, ensure_ascii=False, indent=2))


def cmd_routine_log(args):
    tasks = load_tasks()
    task = None
    for t in tasks:
        if t["id"] == args.task_id and t.get("type") == "routine":
            task = t
            break

    if not task:
        print(json.dumps({"error": f"例行任务 {args.task_id} 未找到"}, ensure_ascii=False))
        sys.exit(1)

    comp_log = task.get("completion_log") or []
    print(json.dumps({
        "task_id": task["id"],
        "title": task["title"],
        "routine_type": task.get("routine_type"),
        "routine_time": task.get("routine_time"),
        "total_completions": len(comp_log),
        "completion_log": comp_log,
    }, ensure_ascii=False, indent=2))


def cmd_routine_pause(args):
    tasks = load_tasks()
    for t in tasks:
        if t["id"] == args.task_id and t.get("type") == "routine":
            if t.get("status") != "active":
                print(json.dumps({"error": f"任务当前状态为 {t.get('status')}，无法暂停"}, ensure_ascii=False))
                sys.exit(1)
            t["status"] = "paused"
            t["updated_at"] = datetime.now().isoformat()
            save_tasks(tasks)
            print(json.dumps({"task_id": t["id"], "title": t["title"], "status": "paused"}, ensure_ascii=False, indent=2))
            return
    print(json.dumps({"error": f"例行任务 {args.task_id} 未找到"}, ensure_ascii=False))
    sys.exit(1)


def cmd_routine_resume(args):
    tasks = load_tasks()
    for t in tasks:
        if t["id"] == args.task_id and t.get("type") == "routine":
            if t.get("status") != "paused":
                print(json.dumps({"error": f"任务当前状态为 {t.get('status')}，无法恢复"}, ensure_ascii=False))
                sys.exit(1)
            t["status"] = "active"
            t["updated_at"] = datetime.now().isoformat()
            save_tasks(tasks)
            print(json.dumps({"task_id": t["id"], "title": t["title"], "status": "active"}, ensure_ascii=False, indent=2))
            return
    print(json.dumps({"error": f"例行任务 {args.task_id} 未找到"}, ensure_ascii=False))
    sys.exit(1)


# ── Module 十三: export ──────────────────────────────────────────────────────

def _check_docx():
    """Check python-docx is installed. Return the module or exit with hint."""
    try:
        import docx
        return docx
    except ImportError:
        print(json.dumps({
            "error": "python-docx 库未安装。请运行: pip install python-docx"
        }, ensure_ascii=False))
        sys.exit(1)


def _generate_export_filename(prefix):
    """Generate a timestamped .docx filename."""
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    return f"{prefix}_{ts}.docx"


def _generate_pdf_filename(prefix):
    """Generate a timestamped .pdf filename."""
    ts = datetime.now().strftime("%Y%m%d")
    return f"{prefix}_{ts}.pdf"


def _check_fpdf2():
    """Check fpdf2 is installed. Return the module or exit with hint."""
    try:
        from fpdf import FPDF
        return FPDF
    except ImportError:
        print(json.dumps({
            "error": "fpdf2 库未安装。请运行: pip install fpdf2"
        }, ensure_ascii=False))
        sys.exit(1)


def _find_chinese_font():
    """Find a Chinese-capable TTF font on the system."""
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/msyhbd.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/System/Library/Fonts/PingFang.ttc",
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return None


def _setup_doc(doc, title_text):
    """Add title, generation date, and footer to a Document."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    title = doc.add_heading(title_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"生成日期: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    # footer
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Generated by Time Planner Skill")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _styled_table(doc, headers, rows, col_widths=None):
    """Create a table with bold header row and grid style."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, w in enumerate(col_widths):
                if ci < len(row.cells):
                    row.cells[ci].width = w
    return table


def _export_plan(args):
    """Export daily plan as a Word document with time-block table."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    tasks = load_tasks()
    profile = load_energy()
    incomplete = [
        t for t in tasks
        if t["status"] not in ("completed", "cancelled")
        and not t.get("has_subtasks", False)
    ]
    sorted_incomplete, cycle_detected = topological_sort(incomplete)
    boss_tasks = find_boss_tasks(sorted_incomplete)
    adjusted_list, total_adjusted, adjustment_active = build_adjusted_task_list(
        sorted_incomplete, profile, sort_by_priority=False
    )
    for entry in adjusted_list:
        _annotate_boss_info(entry, boss_tasks)

    daily_limit_minutes = int(profile.get("daily_limit_hours", 8.0) * 60)
    overload = total_adjusted > daily_limit_minutes

    doc = _setup_doc_base()
    _setup_doc(doc, "📅 每日时间规划")

    # summary info
    info_lines = [
        f"规划日期: {datetime.now().strftime('%Y-%m-%d')}",
        f"时间范围: {args.day_start or '09:00'} - {args.day_end or '18:00'}",
        f"修正后总耗时: {total_adjusted} 分钟 ({total_adjusted / 60:.1f} 小时)",
        f"每日上限: {daily_limit_minutes} 分钟 ({profile.get('daily_limit_hours', 8.0)} 小时)",
    ]
    if adjustment_active:
        info_lines.append("⚠ 已启用精力偏差修正")
    if cycle_detected:
        info_lines.append("⚠ 检测到循环依赖，排序结果可能不完整")
    if overload:
        info_lines.append(f"⚠ 过载预警：总耗时超出每日上限")
    for line in info_lines:
        doc.add_paragraph(line)

    # time block table
    headers = ["时间段", "任务", "优先级", "预估", "修正后", "风险", "备注"]
    rows = []
    for at in adjusted_list:
        remark = ""
        if at.get("is_boss"):
            remark = "🐉 Boss任务"
        elif at.get("blocked_by_boss"):
            remark = f"被Boss阻塞: {at.get('blocked_by_boss_title', '')}"
        if at.get("procrastination_risk") == "high":
            remark += " ⚡ 高风险"
        row = [
            at.get("time", ""),
            at["title"],
            at["priority"],
            f"{at['estimated_minutes']}'",
            f"{at['adjusted_minutes']}'",
            at.get("procrastination_risk", "low"),
            remark,
        ]
        rows.append(row)

    # calculate time slots
    sh, sm = 9, 0
    eh, em = 18, 0
    if args.day_start:
        try:
            sh, sm = map(int, args.day_start.split(":"))
        except (ValueError, AttributeError):
            pass
    if args.day_end:
        try:
            eh, em = map(int, args.day_end.split(":"))
        except (ValueError, AttributeError):
            pass

    cursor = sh * 60 + sm
    day_end = eh * 60 + em
    for i, at in enumerate(adjusted_list):
        dur = at["adjusted_minutes"]
        if cursor + dur > day_end:
            break
        start_str = f"{cursor // 60:02d}:{cursor % 60:02d}"
        end_str = f"{(cursor + dur) // 60:02d}:{(cursor + dur) % 60:02d}"
        rows[i][0] = f"{start_str} - {end_str}"
        cursor += dur

    table = _styled_table(doc, headers, rows)
    # apply bold to Boss task rows
    for ri, row_data in enumerate(rows):
        if _has_boss(row_data):
            for cell in table.rows[ri + 1].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True

    if overload:
        doc.add_paragraph()
        warning = doc.add_paragraph()
        wr = warning.add_run("⚠ 过载提醒：当前计划超出每日上限，建议优先完成排序靠前的任务。")
        wr.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    output = args.output or _generate_export_filename("plan")
    doc.save(output)
    print(json.dumps({"exported": output, "type": "plan"}, ensure_ascii=False))


def _has_boss(row_data):
    """Check if a plan row has Boss marking in remarks."""
    return any("Boss任务" in str(c) for c in row_data)


def _export_list(args):
    """Export task list as a Word document."""
    tasks = load_tasks()
    if args.status:
        tasks = [t for t in tasks if t["status"] == args.status]
    else:
        tasks = [t for t in tasks if t["status"] not in ("completed", "cancelled")]

    tasks.sort(key=lambda t: (PRIORITY_WEIGHT.get(t["priority"], 1), t.get("deadline") or ""), reverse=True)

    doc = _setup_doc_base()
    _setup_doc(doc, "📋 任务清单")

    doc.add_paragraph(f"共 {len(tasks)} 个任务")

    headers = ["标题", "优先级", "状态", "预估(分钟)", "截止时间", "标签", "依赖数"]
    rows = []
    for t in tasks:
        rows.append([
            t["title"],
            t["priority"],
            t["status"],
            str(t.get("estimated_minutes", 30)),
            t.get("deadline") or "—",
            ", ".join(t.get("tags", [])),
            str(len(t.get("dependencies", []))),
        ])
    _styled_table(doc, headers, rows)

    output = args.output or _generate_export_filename("list")
    doc.save(output)
    print(json.dumps({"exported": output, "type": "list", "count": len(tasks)}, ensure_ascii=False))


def _export_stats(args):
    """Export growth statistics report as a Word document."""
    from docx.shared import Pt, RGBColor

    ach_data = load_achievements()
    stats = ach_data["stats"]
    profile = load_energy()

    doc = _setup_doc_base()
    _setup_doc(doc, "📊 成长统计报告")

    # overview
    doc.add_heading("总览", level=2)
    overview_headers = ["指标", "数值"]
    overview_rows = [
        ["累计完成任务", str(stats.get("total_tasks_completed", 0))],
        ["Boss 击败数", str(stats.get("total_bosses_defeated", 0))],
        ["当前连续天数", f"{stats.get('current_streak', 0)} 天"],
        ["最长连续天数", f"{stats.get('longest_streak', 0)} 天"],
        ["过载保护触发次数", str(stats.get("overload_trigger_count", 0))],
        ["成就解锁数", f"{len(ach_data['achievements'])} / {len(ALL_ACHIEVEMENTS)}"],
        ["总复盘次数", str(profile.get("total_reviews", 0))],
    ]
    _styled_table(doc, overview_headers, overview_rows)

    # deviation summary
    doc.add_heading("精力偏差分析", level=2)
    global_rate = profile.get("global_avg_deviation_rate", 0.0)
    sign = "+" if global_rate > 0 else ""
    doc.add_paragraph(f"全局平均偏差率: {sign}{round(global_rate * 100, 1)}%")
    doc.add_paragraph(f"全局平均拖延次数: {profile.get('global_avg_procrastination', 0.0)}")
    doc.add_paragraph(f"全局平均改期次数: {profile.get('global_avg_deadline_adjusted', 0.0)}")

    if profile.get("tag_profiles"):
        doc.add_heading("标签偏差明细", level=3)
        tag_headers = ["标签", "复盘次数", "平均偏差率", "平均拖延", "平均改期"]
        tag_rows = []
        for tag, tp in profile["tag_profiles"].items():
            tag_rows.append([
                tag,
                str(tp.get("count", 0)),
                f"{'+' if tp.get('avg_deviation_rate', 0) > 0 else ''}{round(tp.get('avg_deviation_rate', 0) * 100, 1)}%",
                f"{tp.get('avg_procrastination_count', 0.0):.1f}",
                f"{tp.get('avg_deadline_adjusted', 0.0):.1f}",
            ])
        _styled_table(doc, tag_headers, tag_rows)

    # streak info
    doc.add_heading("连续记录", level=2)
    doc.add_paragraph(
        f"最近完成任务日期: {stats.get('last_completed_date') or '—'}\n"
        f"当前连续: {stats.get('current_streak', 0)} 天 | "
        f"最长连续: {stats.get('longest_streak', 0)} 天"
    )

    output = args.output or _generate_export_filename("stats")
    doc.save(output)
    print(json.dumps({"exported": output, "type": "stats"}, ensure_ascii=False))


def _export_achievements(args):
    """Export achievements list as a Word document."""
    from docx.shared import Pt, RGBColor

    ach_data = load_achievements()
    unlocked_ids = {a["id"] for a in ach_data["achievements"]}
    unlocked_map = {a["id"]: a for a in ach_data["achievements"]}

    doc = _setup_doc_base()
    _setup_doc(doc, "🏆 成就清单")

    unlocked_count = len(unlocked_ids)
    total_count = len(ALL_ACHIEVEMENTS)
    doc.add_paragraph(f"已解锁: {unlocked_count} / {total_count}")

    headers = ["图标", "成就名称", "描述", "状态", "解锁时间"]
    rows = []
    for ach in ALL_ACHIEVEMENTS:
        unlocked = ach["id"] in unlocked_ids
        status_text = "✅ 已解锁" if unlocked else "🔒 未解锁"
        unlocked_at = unlocked_map[ach["id"]].get("unlocked_at", "—") if unlocked else "—"
        rows.append([
            ach["icon"],
            ach["name"],
            ach["description"],
            status_text,
            unlocked_at,
        ])

    table = _styled_table(doc, headers, rows)
    # color unlocked rows green, locked rows gray
    from docx.shared import RGBColor
    for ri, ach in enumerate(ALL_ACHIEVEMENTS):
        unlocked = ach["id"] in unlocked_ids
        color = RGBColor(0x00, 0x80, 0x00) if unlocked else RGBColor(0x99, 0x99, 0x99)
        for cell in table.rows[ri + 1].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = color

    output = args.output or _generate_export_filename("achievements")
    doc.save(output)
    print(json.dumps({"exported": output, "type": "achievements"}, ensure_ascii=False))


def _setup_doc_base():
    """Create a bare Document for export (called before _setup_doc)."""
    from docx import Document
    return Document()


def _export_plan_pdf(args):
    """Export daily plan as a PDF file using fpdf2."""
    _check_fpdf2()
    from fpdf import FPDF

    font_path = _find_chinese_font()
    if font_path is None:
        print(json.dumps({
            "error": "未找到中文字体文件，无法生成 PDF。请安装中文字体或使用 export plan (docx) 代替。"
        }, ensure_ascii=False))
        sys.exit(1)

    tasks = load_tasks()
    profile = load_energy()
    incomplete = [
        t for t in tasks
        if t["status"] not in ("completed", "cancelled")
        and not t.get("has_subtasks", False)
    ]
    sorted_incomplete, cycle_detected = topological_sort(incomplete)
    boss_tasks = find_boss_tasks(sorted_incomplete)
    adjusted_list, total_adjusted, adjustment_active = build_adjusted_task_list(
        sorted_incomplete, profile, sort_by_priority=False
    )
    for entry in adjusted_list:
        _annotate_boss_info(entry, boss_tasks)

    daily_limit_minutes = int(profile.get("daily_limit_hours", 8.0) * 60)
    overload = total_adjusted > daily_limit_minutes

    day_start = args.day_start or "09:00"
    day_end = args.day_end or "18:00"
    try:
        sh, sm = map(int, day_start.split(":"))
        eh, em = map(int, day_end.split(":"))
    except (ValueError, AttributeError):
        sh, sm = 9, 0
        eh, em = 18, 0

    # Create PDF
    pdf = FPDF(orientation='P', unit='mm', format='A4')
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_font('zh', '', font_path)
    pdf.add_font('zh', 'B', font_path)
    pdf.add_page()

    # Title
    pdf.set_font('zh', 'B', 18)
    pdf.cell(0, 12, '📋 今日时间规划', new_x="LMARGIN", new_y="NEXT", align='C')
    pdf.ln(2)

    # Date
    pdf.set_font('zh', '', 9)
    pdf.set_text_color(102, 102, 102)
    pdf.cell(0, 6, f"生成日期: {datetime.now().strftime('%Y-%m-%d %H:%M')}", new_x="LMARGIN", new_y="NEXT", align='C')
    pdf.ln(4)

    # Summary info
    pdf.set_text_color(0, 0, 0)
    pdf.set_font('zh', '', 9)
    info_lines = [
        f"时间范围: {day_start} - {day_end}",
        f"修正后总耗时: {total_adjusted} 分钟 ({total_adjusted / 60:.1f} 小时)",
        f"每日上限: {daily_limit_minutes} 分钟 ({profile.get('daily_limit_hours', 8.0)} 小时)",
    ]
    if adjustment_active:
        info_lines.append("⚠ 已启用精力偏差修正")
    if cycle_detected:
        info_lines.append("⚠ 检测到循环依赖")
    for line in info_lines:
        pdf.cell(0, 5, line, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Overload warning
    if overload:
        pdf.set_font('zh', 'B', 10)
        pdf.set_text_color(255, 0, 0)
        pdf.cell(0, 6, f"⚠ 过载预警：总耗时超出每日上限 {total_adjusted - daily_limit_minutes} 分钟", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(3)

    # Time block table
    col_widths = [35, 65, 18, 28, 24]  # 时间, 任务, 优先级, Boss标记, 修正后
    headers = ["时间", "任务", "优先级", "Boss标记", "修正后"]
    pdf.set_fill_color(240, 240, 240)
    pdf.set_font('zh', 'B', 9)
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        pdf.cell(w, 8, h, border=1, fill=True, align='C')
    pdf.ln()

    cursor = sh * 60 + sm
    day_end_min = eh * 60 + em
    pdf.set_font('zh', '', 9)

    for at in adjusted_list:
        dur = at["adjusted_minutes"]
        if cursor + dur > day_end_min:
            break

        start_str = f"{cursor // 60:02d}:{cursor % 60:02d}"
        end_str = f"{(cursor + dur) // 60:02d}:{(cursor + dur) % 60:02d}"
        time_slot = f"{start_str} - {end_str}"

        is_boss = at.get("is_boss", False)
        boss_mark = "🐉 Boss任务" if is_boss else ("🔒 被阻塞" if at.get("blocked_by_boss") else "-")
        adjusted_str = f"{at['adjusted_minutes']}'"

        is_bold = is_boss or overload
        font_style = 'B' if is_bold else ''
        pdf.set_font('zh', font_style, 9)

        row_data = [time_slot, at["title"], at["priority"], boss_mark, adjusted_str]
        for val, w in zip(row_data, col_widths):
            if overload and not is_boss:
                pdf.set_text_color(255, 0, 0)
            pdf.cell(w, 7, str(val)[:30], border=1, align='C' if val != at["title"] else 'L')
            pdf.set_text_color(0, 0, 0)
        pdf.ln()
        cursor += dur

    # Output
    output = args.output or _generate_pdf_filename("plan")
    if not output.endswith(".pdf"):
        output += ".pdf"
    pdf.output(output)
    print(json.dumps({"exported": output, "type": "plan_pdf"}, ensure_ascii=False))


def cmd_export(args):
    """Export data to a Word (.docx) file or PDF."""
    # PDF path for plan type
    if getattr(args, "pdf", False):
        if args.type != "plan":
            print(json.dumps({
                "error": "--pdf 目前仅支持 plan 类型，其他类型请使用 docx 导出后再转换"
            }, ensure_ascii=False))
            sys.exit(1)
        _export_plan_pdf(args)
        return

    _check_docx()

    if args.type == "plan":
        _export_plan(args)
    elif args.type == "list":
        _export_list(args)
    elif args.type == "stats":
        _export_stats(args)
    elif args.type == "achievements":
        _export_achievements(args)
    else:
        print(json.dumps({"error": f"未知导出类型: {args.type}，可选: plan, list, stats, achievements"}, ensure_ascii=False))
        sys.exit(1)


# ── cli ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Time Planner task manager")
    sub = parser.add_subparsers(dest="command", required=True)

    # add
    p_add = sub.add_parser("add", help="Add a new task")
    p_add.add_argument("--title", required=True)
    p_add.add_argument("--description", default="")
    p_add.add_argument("--priority", default="medium", choices=VALID_PRIORITIES)
    p_add.add_argument("--estimated-minutes", type=int, default=30)
    p_add.add_argument("--deadline", default=None)
    p_add.add_argument("--tags", default=None)

    # list
    p_list = sub.add_parser("list", help="List tasks")
    p_list.add_argument("--status", choices=VALID_STATUSES)
    p_list.add_argument("--priority", choices=VALID_PRIORITIES)

    # modify
    p_mod = sub.add_parser("modify", help="Modify a task")
    p_mod.add_argument("task_id")
    p_mod.add_argument("--title")
    p_mod.add_argument("--description")
    p_mod.add_argument("--priority", choices=VALID_PRIORITIES)
    p_mod.add_argument("--status", choices=VALID_STATUSES)
    p_mod.add_argument("--estimated-minutes", type=int)
    p_mod.add_argument("--deadline", default=argparse.SUPPRESS,
                       help="Set deadline (ISO format, e.g. 2026-05-16T18:00)")
    p_mod.add_argument("--clear-deadline", dest="deadline_flag", action="store_true",
                       help="Remove the deadline")
    p_mod.add_argument("--tags")

    # delete
    p_del = sub.add_parser("delete", help="Delete a task")
    p_del.add_argument("task_id")

    # plan
    p_plan = sub.add_parser("plan", help="Generate daily time plan")
    p_plan.add_argument("--day-start", default="09:00")
    p_plan.add_argument("--day-end", default="18:00")
    p_plan.add_argument("--force", action="store_true")
    p_plan.add_argument("--dynamic", action="store_true", help="Add next_best suggestion to output")
    p_plan.add_argument("--quick", action="store_true", help="Skip overload confirmation, auto-force if overloaded")

    # review
    p_rev = sub.add_parser("review", help="Record actual time and update energy profile")
    p_rev.add_argument("task_id")
    p_rev.add_argument("actual_minutes", type=int, help="Actual minutes spent")

    # energy
    sub.add_parser("energy", help="Show energy deviation profile summary")

    # limit
    p_lim = sub.add_parser("limit", help="Show or set daily working hour limit")
    p_lim.add_argument("--set", type=float, dest="set", default=None, help="Set daily limit in hours")

    # deps
    p_deps = sub.add_parser("deps", help="View / add / remove task dependencies")
    p_deps.add_argument("task_id", help="Task ID")
    p_deps.add_argument("--add", default=None, help="Add a dependency (task ID)")
    p_deps.add_argument("--remove", default=None, help="Remove a dependency (task ID)")

    # config
    p_cfg = sub.add_parser("config", help="Show or set config (style, data_dir, etc.)")
    p_cfg.add_argument("--set-style", dest="set_style", default=None, choices=VALID_STYLES,
                       help="Set conversation style")
    p_cfg.add_argument("--set-data-dir", dest="set_data_dir", default=None,
                       help="Set custom data directory for tasks.json, energy_profile.json, achievements.json")
    p_cfg.add_argument("--show-data-dir", dest="show_data_dir", action="store_true",
                       help="Show current data directory")
    p_cfg.add_argument("--reset-data-dir", dest="reset_data_dir", action="store_true",
                       help="Reset data directory to default (skill folder)")

    # predict (Module 十)
    p_pred = sub.add_parser("predict", help="Predict task duration and procrastination risk")
    p_pred.add_argument("task_id", help="Task ID to predict")

    # warn (Module 十)
    sub.add_parser("warn", help="Scan for high procrastination risk tasks")

    # break (Module 十一)
    p_break = sub.add_parser("break", help="Split a large task into sub-tasks by phases")
    p_break.add_argument("task_id", help="Task ID to split")

    # stats (Module 十二)
    sub.add_parser("stats", help="Show growth tracking statistics")

    # achievements (Module 十二)
    sub.add_parser("achievements", help="List all achievements with unlock status")

    # export (Module 十三)
    p_export = sub.add_parser("export", help="Export data to Word (.docx) file")
    p_export.add_argument("type", choices=["plan", "list", "stats", "achievements"],
                          help="导出类型: plan, list, stats, achievements")
    p_export.add_argument("--output", default=None, help="输出文件路径 (默认自动生成时间戳文件名)")
    p_export.add_argument("--status", choices=VALID_STATUSES, default=None,
                          help="(仅 list) 按状态筛选任务")
    p_export.add_argument("--day-start", default=None, help="(仅 plan) 开始时间，如 09:00")
    p_export.add_argument("--day-end", default=None, help="(仅 plan) 结束时间，如 18:00")
    p_export.add_argument("--pdf", action="store_true", dest="pdf", default=False,
                          help="(仅 plan) 导出为 PDF 文件（需要 fpdf2 库）")

    # routine (Module 十四)
    p_routine = sub.add_parser("routine", help="Manage routine/recurring tasks")
    p_routine.add_argument("action", choices=["add", "list", "done", "log", "pause", "resume"],
                           help="routine 子命令: add, list, done, log, pause, resume")
    p_routine.add_argument("--title", default=None, help="Routine task title")
    p_routine.add_argument("--description", default=None)
    p_routine.add_argument("--routine-type", dest="routine_type", choices=sorted(VALID_ROUTINE_TYPES),
                           default=None, help="Routine frequency: daily, weekly, monthly")
    p_routine.add_argument("--routine-time", dest="routine_time", default=None,
                           help="Suggested execution time (HH:MM)")
    p_routine.add_argument("--routine-days", dest="routine_days", default=None,
                           help="For weekly: comma-separated weekday numbers (1=Mon..7=Sun)")
    p_routine.add_argument("--estimated-minutes", dest="estimated_minutes", type=int, default=None)
    p_routine.add_argument("--priority", default="medium", choices=VALID_PRIORITIES)
    p_routine.add_argument("--tags", default=None)
    p_routine.add_argument("--task-id", dest="task_id", default=None, help="Routine task ID")

    args = parser.parse_args()

    cmds = {
        "add": cmd_add,
        "list": cmd_list,
        "modify": cmd_modify,
        "delete": cmd_delete,
        "plan": cmd_plan,
        "review": cmd_review,
        "energy": cmd_energy,
        "limit": cmd_limit,
        "deps": cmd_deps,
        "config": cmd_config,
        "predict": cmd_predict,
        "warn": cmd_warn,
        "break": cmd_break,
        "stats": cmd_stats,
        "achievements": cmd_achievements,
        "export": cmd_export,
        "routine": cmd_routine,
    }
    cmds[args.command](args)


if __name__ == "__main__":
    main()
